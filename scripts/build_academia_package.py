"""Build the Academia.edu technical report, visual appendix, and figures."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
ACADEMIA = ROOT / "academia"
FIGURES = ACADEMIA / "figures"
OUTPUT = ACADEMIA / "output"
PROFILE = ROOT / "assets" / "emmanuel-onuoha.jpg"
NAVY = "#173B67"
ORANGE = "#F28E2B"
BLUE = "#3B7FB6"
LIGHT_BLUE = "#DCE8F5"


def configure_plot() -> None:
    plt.rcParams.update(
        {
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#B7C6D9",
            "axes.labelcolor": NAVY,
            "axes.titlecolor": NAVY,
            "axes.grid": True,
            "grid.color": "#E1E9F2",
            "grid.linewidth": 0.8,
            "font.family": "DejaVu Sans",
            "font.size": 10,
        }
    )


def save_figure(fig: plt.Figure, name: str) -> None:
    fig.savefig(FIGURES / name, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_figures() -> None:
    FIGURES.mkdir(parents=True, exist_ok=True)
    configure_plot()

    labels = [
        "Months 1–6 vs 49–72",
        "Month-to-month vs two-year",
        "San Diego vs elsewhere",
        "Without dependents",
        "Fibre optic vs DSL",
        "Bank withdrawal vs credit card",
        "Senior vs non-senior",
        "Without online security",
        "Without premium support",
    ]
    estimates = np.array([43.82, 43.30, 39.99, 26.04, 22.14, 19.52, 18.08, 16.72, 16.02])
    lows = np.array([40.96, 41.48, 34.19, 24.24, 19.53, 17.52, 15.04, 14.67, 13.96])
    highs = np.array([46.61, 45.04, 45.40, 27.71, 24.66, 21.48, 21.15, 18.68, 18.00])
    order = np.arange(len(labels))[::-1]
    fig, ax = plt.subplots(figsize=(9.0, 5.8))
    ax.errorbar(
        estimates[order], order,
        xerr=np.vstack((estimates[order] - lows[order], highs[order] - estimates[order])),
        fmt="o", color=BLUE, ecolor="#89A7C7", capsize=4, markersize=7,
    )
    ax.axvline(0, color="#566573", linewidth=1)
    ax.set_yticks(order, [labels[i] for i in order])
    ax.set_xlabel("Absolute churn-risk difference (percentage points)")
    ax.set_title("Selected Validated Churn-Risk Differences", loc="left", fontsize=16, fontweight="bold")
    ax.text(0, 1.015, "Point estimates with 95% confidence intervals; observational comparisons", transform=ax.transAxes, color="#53657A")
    ax.grid(axis="y", visible=False)
    save_figure(fig, "01_validated_risk_differences.png")

    shutil.copy2(ROOT / "artifacts" / "figures" / "cross-validated-precision-recall-tradeoff.png", FIGURES / "02_cross_validated_precision_recall.png")

    metrics = pd.read_json(ROOT / "artifacts" / "evaluation" / "selected_threshold_holdout_metrics.json", typ="series")
    matrix = np.array([[metrics["true_negatives"], metrics["false_positives"]], [metrics["false_negatives"], metrics["true_positives"]]])
    fig, ax = plt.subplots(figsize=(6.4, 5.4))
    im = ax.imshow(matrix, cmap="Blues")
    for (row, col), value in np.ndenumerate(matrix):
        ax.text(col, row, f"{int(value):,}", ha="center", va="center", fontsize=18, fontweight="bold", color="white" if value > 450 else NAVY)
    ax.set_xticks([0, 1], ["Predicted retain", "Predicted churn"])
    ax.set_yticks([0, 1], ["Actual retain", "Actual churn"])
    ax.set_title("Holdout Confusion Matrix at Threshold 0.32", loc="left", fontsize=15, fontweight="bold")
    ax.grid(False)
    fig.colorbar(im, ax=ax, fraction=0.045, pad=0.04)
    save_figure(fig, "03_confusion_matrix.png")

    shutil.copy2(ROOT / "artifacts" / "figures" / "why-customers-churned.png", FIGURES / "04_why_customers_churned.png")

    segments = pd.read_csv(ROOT / "data" / "processed" / "prioritized_risk_segments.csv").head(7).copy()
    segments["label"] = segments["segment"].replace(
        {
            "Month-to-month without premium support": "Month-to-month, no support",
            "Fiber optic on month-to-month contract": "Fibre, month-to-month",
            "High-value with high descriptive risk": "High-value, high risk",
            "Early-tenure month-to-month": "Early tenure, month-to-month",
            "Early-tenure fiber month-to-month": "Early fibre, month-to-month",
            "Senior with $90+ monthly charge": "Senior, $90+ charge",
            "San Diego customers": "San Diego",
        }
    )
    fig, ax = plt.subplots(figsize=(9.2, 6.0))
    sizes = 70 + 500 * segments["active_monthly_charge_exposure"] / segments["active_monthly_charge_exposure"].max()
    scatter = ax.scatter(segments["active_customers"], segments["observed_churn_rate"] * 100, s=sizes, c=segments["active_monthly_charge_exposure"], cmap="Blues", edgecolor="white", linewidth=1.2)
    for _, row in segments.iterrows():
        ax.annotate(row["label"], (row["active_customers"], row["observed_churn_rate"] * 100), xytext=(5, 5), textcoords="offset points", fontsize=8)
    ax.set_xlabel("Active customers")
    ax.set_ylabel("Observed churn rate")
    ax.yaxis.set_major_formatter(lambda x, pos: f"{x:.0f}%")
    ax.set_title("Prioritised Customer-Risk Segments", loc="left", fontsize=16, fontweight="bold")
    cbar = fig.colorbar(scatter, ax=ax, pad=0.02)
    cbar.set_label("Active monthly-charge exposure ($)")
    save_figure(fig, "05_priority_segments.png")

    scenarios = pd.read_csv(ROOT / "artifacts" / "retention" / "retention_scenario_results.csv")
    x = np.arange(len(scenarios))
    width = 0.34
    fig, ax = plt.subplots(figsize=(8.6, 5.5))
    ax.bar(x - width / 2, scenarios["total_campaign_cost"], width, label="Campaign cost", color="#9CB9D6")
    ax.bar(x + width / 2, scenarios["expected_retained_gross_margin"], width, label="Retained gross margin", color=BLUE)
    for index, value in enumerate(scenarios["estimated_net_benefit"]):
        ax.text(index, max(scenarios.loc[index, "total_campaign_cost"], scenarios.loc[index, "expected_retained_gross_margin"]) + 3500, f"Net: ${value:,.0f}", ha="center", color=ORANGE if value >= 0 else "#B33A3A", fontweight="bold")
    ax.set_xticks(x, scenarios["scenario"])
    ax.set_ylabel("Estimated value ($)")
    ax.set_title("Retention Scenario Economics", loc="left", fontsize=16, fontweight="bold")
    ax.legend(frameon=False)
    ax.yaxis.set_major_formatter(lambda value, pos: f"${value/1000:.0f}k")
    save_figure(fig, "06_retention_scenarios.png")

    shutil.copy2(ROOT / "artifacts" / "testing" / "screenshots" / "dark-decision-centre-1440x1000.png", FIGURES / "07_decision_centre.png")


def clean_inline(text: str) -> str:
    text = re.sub(r"(https?://[^\s<]+)", r"<link href='\1' color='#285F97'>\1</link>", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", text)
    return text


def markdown_blocks(path: Path) -> list[tuple[str, object]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, object]] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(item.strip() for item in paragraph)))
            paragraph.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
        elif stripped == "{{PAGEBREAK}}":
            flush(); blocks.append(("pagebreak", ""))
        elif stripped == "{{PROFILE}}":
            flush(); blocks.append(("profile", PROFILE))
        elif stripped.startswith("!["):
            flush()
            match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
            if match:
                blocks.append(("figure", (match.group(1), ACADEMIA / match.group(2))))
        elif stripped.startswith("### "):
            flush(); blocks.append(("heading3", stripped[4:]))
        elif stripped.startswith("## "):
            flush(); blocks.append(("heading2", stripped[3:]))
        elif stripped.startswith("# "):
            flush(); blocks.append(("heading1", stripped[2:]))
        elif stripped.startswith("- "):
            flush(); blocks.append(("bullet", stripped[2:]))
        elif re.match(r"^\d+\. ", stripped):
            flush(); blocks.append(("number", re.sub(r"^\d+\. ", "", stripped)))
        elif stripped.startswith("> "):
            flush(); blocks.append(("quote", stripped[2:]))
        elif stripped == "---":
            flush(); blocks.append(("rule", ""))
        else:
            paragraph.append(stripped)
    flush()
    return blocks


def add_docx_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = field
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_docx_cell_border(section) -> None:
    header = section.header.paragraphs[0]
    header.text = "TELECOM CUSTOMER CHURN INTELLIGENCE  |  EMMANUEL ONUOHA"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(23, 59, 103)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Developed and Designed by Emmanuel Onuoha  •  ")
    add_docx_field(footer, "PAGE")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(80, 95, 115)


def build_docx(source: Path, destination: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78); section.right_margin = Inches(0.78)
    set_docx_cell_border(section)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for style_name, size, color in [("Title", 25, RGBColor(23, 59, 103)), ("Heading 1", 17, RGBColor(23, 59, 103)), ("Heading 2", 13, RGBColor(23, 59, 103)), ("Heading 3", 11, RGBColor(242, 142, 43))]:
        styles[style_name].font.name = "Aptos Display"; styles[style_name].font.size = Pt(size); styles[style_name].font.color.rgb = color

    first_heading = True
    for kind, value in markdown_blocks(source):
        if kind == "heading1":
            p = doc.add_paragraph(style="Title" if first_heading else "Heading 1"); first_heading = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if p.style.name == "Title" else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(str(value))
        elif kind in {"heading2", "heading3"}:
            doc.add_paragraph(str(value), style="Heading 1" if kind == "heading2" else "Heading 2")
        elif kind == "paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Word receives clean readable text; basic Markdown marks are removed.
            p.add_run(re.sub(r"[*`]", "", str(value)).replace("  ", " "))
        elif kind in {"bullet", "number"}:
            doc.add_paragraph(re.sub(r"[*`]", "", str(value)), style="List Bullet" if kind == "bullet" else "List Number")
        elif kind == "quote":
            p = doc.add_paragraph(re.sub(r"[*`]", "", str(value)))
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs: run.font.italic = True; run.font.color.rgb = RGBColor(80, 95, 115)
        elif kind == "profile":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(value), width=Inches(1.35))
        elif kind == "figure":
            caption, image_path = value
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image_path), width=Inches(6.55))
            cap = doc.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs: run.font.size = Pt(8.5); run.font.italic = True; run.font.color.rgb = RGBColor(70, 85, 105)
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "rule":
            doc.add_paragraph("—" * 18).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.core_properties.title = source.stem.replace("_", " ")
    doc.core_properties.author = "Emmanuel Onuoha"
    doc.core_properties.subject = "Telecom customer churn, explainable modelling, and prescriptive retention analytics"
    doc.save(destination)


class NumberedDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs):
        super().__init__(filename, **kwargs)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="normal")
        self.addPageTemplates(PageTemplate(id="main", frames=frame, onPage=self.draw_header_footer))

    @staticmethod
    def draw_header_footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#DCE8F5")); canvas.line(doc.leftMargin, A4[1] - 38, A4[0] - doc.rightMargin, A4[1] - 38)
        canvas.setFillColor(colors.HexColor(NAVY)); canvas.setFont("Helvetica", 7.5)
        canvas.drawString(doc.leftMargin, A4[1] - 30, "TELECOM CUSTOMER CHURN INTELLIGENCE")
        canvas.setFillColor(colors.HexColor("#53657A")); canvas.drawRightString(A4[0] - doc.rightMargin, A4[1] - 30, "EMMANUEL ONUOHA")
        canvas.setStrokeColor(colors.HexColor("#DCE8F5")); canvas.line(doc.leftMargin, 34, A4[0] - doc.rightMargin, 34)
        canvas.setFont("Helvetica", 7.5); canvas.drawString(doc.leftMargin, 22, "Developed and Designed by Emmanuel Onuoha")
        canvas.drawRightString(A4[0] - doc.rightMargin, 22, f"Page {doc.page}")
        canvas.restoreState()


def build_pdf(source: Path, destination: Path) -> None:
    sample = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("TitleCustom", parent=sample["Title"], fontName="Helvetica-Bold", fontSize=23, leading=27, textColor=colors.HexColor(NAVY), alignment=TA_CENTER, spaceAfter=15),
        "h1": ParagraphStyle("H1Custom", parent=sample["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor(NAVY), spaceBefore=13, spaceAfter=7),
        "h2": ParagraphStyle("H2Custom", parent=sample["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor(NAVY), spaceBefore=9, spaceAfter=5),
        "body": ParagraphStyle("BodyCustom", parent=sample["BodyText"], fontName="Helvetica", fontSize=9.3, leading=13, textColor=colors.HexColor("#25364A"), alignment=TA_JUSTIFY, spaceAfter=6),
        "bullet": ParagraphStyle("BulletCustom", parent=sample["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12.5, leftIndent=14, firstLineIndent=-8, textColor=colors.HexColor("#25364A"), spaceAfter=3),
        "quote": ParagraphStyle("QuoteCustom", parent=sample["BodyText"], fontName="Helvetica-Oblique", fontSize=8.7, leading=12, leftIndent=18, rightIndent=18, textColor=colors.HexColor("#53657A"), backColor=colors.HexColor("#F3F7FB"), borderPadding=7, spaceAfter=8),
        "caption": ParagraphStyle("CaptionCustom", parent=sample["BodyText"], fontName="Helvetica-Oblique", fontSize=8, leading=10, alignment=TA_CENTER, textColor=colors.HexColor("#53657A"), spaceAfter=10),
    }
    story = []
    first_heading = True
    for kind, value in markdown_blocks(source):
        if kind == "heading1":
            story.append(Paragraph(clean_inline(str(value)), styles["title"] if first_heading else styles["h1"])); first_heading = False
        elif kind == "heading2": story.append(Paragraph(clean_inline(str(value)), styles["h1"]))
        elif kind == "heading3": story.append(Paragraph(clean_inline(str(value)), styles["h2"]))
        elif kind == "paragraph": story.append(Paragraph(clean_inline(str(value)).replace("  ", "<br/>"), styles["body"]))
        elif kind == "bullet": story.append(Paragraph("• " + clean_inline(str(value)), styles["bullet"]))
        elif kind == "number": story.append(Paragraph(clean_inline(str(value)), styles["bullet"]))
        elif kind == "quote": story.append(Paragraph(clean_inline(str(value)), styles["quote"]))
        elif kind == "profile":
            image = Image(str(value), width=1.25 * inch, height=1.55 * inch)
            image.hAlign = "CENTER"; story.extend([image, Spacer(1, 7)])
        elif kind == "figure":
            caption, image_path = value
            from PIL import Image as PILImage
            with PILImage.open(image_path) as opened:
                ratio = opened.height / opened.width
            width = 6.55 * inch; height = min(width * ratio, 6.8 * inch)
            if height == 6.8 * inch: width = height / ratio
            image = Image(str(image_path), width=width, height=height); image.hAlign = "CENTER"
            story.append(KeepTogether([image, Spacer(1, 3), Paragraph(clean_inline(caption), styles["caption"])]))
        elif kind == "pagebreak": story.append(PageBreak())
        elif kind == "rule": story.append(Spacer(1, 8))
    doc = NumberedDocTemplate(str(destination), pagesize=A4, rightMargin=48, leftMargin=48, topMargin=52, bottomMargin=45, title=source.stem, author="Emmanuel Onuoha")
    doc.build(story)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_figures()
    publications = [
        (ACADEMIA / "PAPER.md", "Emmanuel_Onuoha_Telecom_Churn_Technical_Report"),
        (ACADEMIA / "VISUAL_APPENDIX.md", "Emmanuel_Onuoha_Telecom_Churn_Visual_Appendix"),
    ]
    for source, stem in publications:
        build_docx(source, OUTPUT / f"{stem}.docx")
        build_pdf(source, OUTPUT / f"{stem}.pdf")
    print(f"Built {len(publications)} publication documents in {OUTPUT}")


if __name__ == "__main__":
    main()
