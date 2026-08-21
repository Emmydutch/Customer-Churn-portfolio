"""Validate the generated Academia publication package and render evidence pages."""

from __future__ import annotations

import json
from pathlib import Path

import pymupdf
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
ACADEMIA = ROOT / "academia"
OUTPUT = ACADEMIA / "output"
PREVIEW = ACADEMIA / "preview"


def validate_pdf(path: Path, minimum_pages: int, phrases: list[str]) -> dict[str, object]:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    missing = [phrase for phrase in phrases if phrase not in text]
    if len(reader.pages) < minimum_pages or missing:
        raise AssertionError(f"PDF validation failed for {path.name}: pages={len(reader.pages)}, missing={missing}")
    return {"file": path.name, "pages": len(reader.pages), "characters": len(text), "size_bytes": path.stat().st_size}


def validate_docx(path: Path, minimum_paragraphs: int, phrases: list[str]) -> dict[str, object]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    missing = [phrase for phrase in phrases if phrase not in text]
    if len(document.paragraphs) < minimum_paragraphs or missing:
        raise AssertionError(f"DOCX validation failed for {path.name}: paragraphs={len(document.paragraphs)}, missing={missing}")
    return {"file": path.name, "paragraphs": len(document.paragraphs), "characters": len(text), "size_bytes": path.stat().st_size}


def render_preview(pdf_path: Path, pages: list[int]) -> list[str]:
    PREVIEW.mkdir(parents=True, exist_ok=True)
    document = pymupdf.open(pdf_path)
    rendered = []
    for page_number in pages:
        if page_number >= len(document):
            continue
        output = PREVIEW / f"{pdf_path.stem}-page-{page_number + 1}.png"
        pixmap = document[page_number].get_pixmap(matrix=pymupdf.Matrix(1.35, 1.35), alpha=False)
        pixmap.save(output)
        rendered.append(output.name)
    document.close()
    return rendered


def main() -> None:
    common = ["Emmanuel Onuoha", "Developed and Designed by Emmanuel Onuoha"]
    report_stem = "Emmanuel_Onuoha_Telecom_Churn_Technical_Report"
    appendix_stem = "Emmanuel_Onuoha_Telecom_Churn_Visual_Appendix"
    results = {
        "pdf": [
            validate_pdf(OUTPUT / f"{report_stem}.pdf", 10, common + ["Abstract", "Limitations and Responsible Use", "$18,775"]),
            validate_pdf(OUTPUT / f"{appendix_stem}.pdf", 6, common + ["Visual Statistics Appendix"]),
        ],
        "docx": [
            validate_docx(OUTPUT / f"{report_stem}.docx", 75, common + ["Abstract", "Limitations and Responsible Use"]),
            validate_docx(OUTPUT / f"{appendix_stem}.docx", 15, common + ["Visual Statistics Appendix"]),
        ],
    }
    results["preview"] = render_preview(OUTPUT / f"{report_stem}.pdf", [0, 2, 6, 10])
    results["status"] = "passed"
    (OUTPUT / "validation_report.json").write_text(json.dumps(results, indent=2), encoding="utf-8")
    print(json.dumps(results, indent=2))


if __name__ == "__main__":
    main()
